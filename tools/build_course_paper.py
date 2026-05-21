from pathlib import Path
from math import atan, sqrt

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
PROJECT = ROOT / "Project"
FIG_DIR = PROJECT / "论文插图"
OUT = PROJECT / "宋佳骏课程论文.docx"


FONT_SONG = "SimSun"
FONT_HEI = "SimHei"
FONT_KAI = "KaiTi"
FONT_EN = "Times New Roman"


def font_path(name: str) -> str:
    candidates = {
        "song": [Path(r"C:\Windows\Fonts\simsun.ttc"), Path(r"C:\Windows\Fonts\msyh.ttc")],
        "hei": [Path(r"C:\Windows\Fonts\simhei.ttf"), Path(r"C:\Windows\Fonts\msyh.ttc")],
        "kai": [Path(r"C:\Windows\Fonts\simkai.ttf"), Path(r"C:\Windows\Fonts\simsun.ttc")],
    }
    for path in candidates[name]:
        if path.exists():
            return str(path)
    return str(Path(r"C:\Windows\Fonts\msyh.ttc"))


def pil_font(kind: str, size: int):
    return ImageFont.truetype(font_path(kind), size)


def draw_round_box(draw, xy, text, font, fill="#FFFFFF", outline="#4A6FA5", text_fill="#111111", radius=14):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=8, align="center")
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text(((x1 + x2 - tw) / 2, (y1 + y2 - th) / 2), text, font=font, fill=text_fill, spacing=8, align="center")


def arrow(draw, p1, p2, text=None, font=None, fill="#4A6FA5"):
    draw.line([p1, p2], fill=fill, width=4)
    x1, y1 = p1
    x2, y2 = p2
    dx = x2 - x1
    dy = y2 - y1
    if abs(dx) >= abs(dy):
        s = 1 if dx >= 0 else -1
        head = [(x2, y2), (x2 - 18 * s, y2 - 10), (x2 - 18 * s, y2 + 10)]
    else:
        s = 1 if dy >= 0 else -1
        head = [(x2, y2), (x2 - 10, y2 - 18 * s), (x2 + 10, y2 - 18 * s)]
    draw.polygon(head, fill=fill)
    if text and font:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        tb = draw.textbbox((0, 0), text, font=font)
        draw.rectangle((mx - (tb[2]-tb[0])/2 - 8, my - 18, mx + (tb[2]-tb[0])/2 + 8, my + 18), fill="#FFFFFF")
        draw.text((mx - (tb[2]-tb[0])/2, my - 12), text, font=font, fill="#2C4D73")


def create_figures():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    title_font = pil_font("hei", 34)
    box_font = pil_font("hei", 28)
    small_font = pil_font("song", 22)
    tiny_font = pil_font("song", 18)

    # Figure 1: system architecture
    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.text((420, 35), "智能坐姿与学习环境监测系统总体架构", font=title_font, fill="#123B63")
    draw_round_box(draw, (610, 320, 890, 500), "STM32F103C8T6\n主控单元", box_font, fill="#EAF2FB")
    draw_round_box(draw, (90, 150, 380, 270), "MPU6050\n姿态传感器", box_font, fill="#F7FBFF")
    draw_round_box(draw, (90, 365, 380, 485), "光敏电阻模块\nAO 模拟输出", box_font, fill="#F7FBFF")
    draw_round_box(draw, (90, 585, 380, 705), "旋转编码器\nA/B/SW", box_font, fill="#F7FBFF")
    draw_round_box(draw, (1100, 145, 1380, 265), "OLED 显示\nAngle / Light", box_font, fill="#F7FBFF")
    draw_round_box(draw, (1100, 365, 1380, 485), "有源蜂鸣器\n坐姿报警", box_font, fill="#F7FBFF")
    draw_round_box(draw, (1100, 585, 1380, 705), "警示 LED\n光照不足提醒", box_font, fill="#F7FBFF")
    arrow(draw, (380, 210), (610, 360), "I2C PB10/PB11", tiny_font)
    arrow(draw, (380, 425), (610, 420), "ADC PA0", tiny_font)
    arrow(draw, (380, 645), (610, 480), "EXTI PA1/PA2, PA3", tiny_font)
    arrow(draw, (890, 360), (1100, 205), "I2C PB8/PB9", tiny_font)
    arrow(draw, (890, 420), (1100, 425), "GPIO PB12", tiny_font)
    arrow(draw, (890, 480), (1100, 645), "GPIO PB13", tiny_font)
    fig1 = FIG_DIR / "图1_系统总体架构.png"
    img.save(fig1)

    # Figure 2: FSM
    img = Image.new("RGB", (1500, 980), "white")
    draw = ImageDraw.Draw(img)
    draw.text((510, 35), "软件三状态菜单有限状态机", font=title_font, fill="#123B63")
    draw_round_box(draw, (590, 120, 910, 220), "系统初始化", box_font, fill="#EEF5FB")
    draw_round_box(draw, (540, 300, 960, 420), "STATE_MONITOR\n主监控界面", box_font, fill="#EAF2FB")
    draw_round_box(draw, (90, 610, 510, 745), "STATE_SET_ANGLE\n角度阈值设置\n编码器调节 0-90°", box_font, fill="#F7FBFF")
    draw_round_box(draw, (990, 610, 1410, 745), "STATE_SET_DELAY\n报警延迟设置\n编码器调节 1-10 s", box_font, fill="#F7FBFF")
    arrow(draw, (750, 220), (750, 300), "上电", tiny_font)
    arrow(draw, (540, 360), (510, 660), "PA3 按下", tiny_font)
    arrow(draw, (510, 680), (990, 680), "PA3 按下", tiny_font)
    arrow(draw, (1200, 610), (960, 380), "PA3 按下", tiny_font)
    draw.line([(730, 420), (730, 500), (835, 500), (835, 420)], fill="#4A6FA5", width=4)
    draw.polygon([(835, 420), (825, 438), (845, 438)], fill="#4A6FA5")
    draw.text((790, 460), "采集-计算-显示-报警循环", font=tiny_font, fill="#2C4D73")
    fig2 = FIG_DIR / "图2_软件状态机.png"
    img.save(fig2)

    # Figure 3: pin wiring layout
    img = Image.new("RGB", (1500, 920), "white")
    draw = ImageDraw.Draw(img)
    draw.text((520, 35), "硬件模块连接与引脚分配示意", font=title_font, fill="#123B63")
    draw_round_box(draw, (560, 180, 940, 700), "STM32F103C8T6\n最小系统板\n\nPB8  PB9\nPB10 PB11\nPA0 PA1 PA2 PA3\nPB12 PB13", box_font, fill="#EAF2FB")
    modules = [
        ((80, 120, 390, 220), "OLED\nSCL-PB8  SDA-PB9", (560, 240)),
        ((80, 310, 390, 410), "MPU6050\nSCL-PB10 SDA-PB11", (560, 330)),
        ((80, 500, 390, 600), "光敏电阻\nAO-PA0", (560, 455)),
        ((1110, 120, 1420, 220), "旋转编码器\nA-PA1 B-PA2 SW-PA3", (940, 485)),
        ((1110, 350, 1420, 450), "蜂鸣器\nIN-PB12 低有效", (940, 570)),
        ((1110, 580, 1420, 680), "LED\nIN-PB13 低有效", (940, 650)),
    ]
    for xy, txt, end in modules:
        draw_round_box(draw, xy, txt, small_font, fill="#F7FBFF")
        start = (xy[2], (xy[1] + xy[3]) // 2) if xy[0] < 750 else (xy[0], (xy[1] + xy[3]) // 2)
        arrow(draw, start, end, None, tiny_font)
    fig3 = FIG_DIR / "图3_硬件连接示意.png"
    img.save(fig3)

    return fig1, fig2, fig3


def set_run_font(run, size=12, name=FONT_SONG, east_asia=None, bold=None, italic=None, color=None):
    east_asia = east_asia or name
    run.font.name = FONT_EN if name == FONT_SONG else name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(p, first_line=True, align=None, before=0, after=0, line=1.0):
    pf = p.paragraph_format
    pf.first_line_indent = Pt(24) if first_line else None
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align


def add_para(doc, text="", size=12, bold=False, first_line=True, align=None, style=None, after=0, before=0):
    p = doc.add_paragraph(style=style)
    set_paragraph_format(p, first_line=first_line, align=align, before=before, after=after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, before=10 if level == 1 else 6, after=6, line=1.0)
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 14 if level == 2 else 12, name=FONT_HEI, east_asia=FONT_HEI, bold=True)
    return p


def set_cell_text(cell, text, size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r, size=12, bold=True)


def add_source(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=10.5, name=FONT_KAI, east_asia=FONT_KAI)


def add_table(doc, caption, headers, rows, widths_cm=None):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        shade_cell(hdr[i], "E8EEF5")
        set_cell_margins(hdr[i])
        if widths_cm:
            hdr[i].width = Cm(widths_cm[i])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], str(text), align=WD_ALIGN_PARAGRAPH.LEFT if len(str(text)) > 8 else WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_margins(cells[i])
            if widths_cm:
                cells[i].width = Cm(widths_cm[i])
    add_source(doc, "资料来源：作者根据系统设计参数整理。")
    return table


def create_document():
    fig1, fig2, fig3 = create_figures()

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SONG)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    normal.font.size = Pt(12)

    # Cover page
    for _ in range(2):
        doc.add_paragraph()
    p = add_para(doc, "辽宁大学", size=22, bold=True, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEI)
    doc.add_paragraph()
    add_para(doc, "本 科 课 程 论 文", size=20, bold=True, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "（2025-2026 学年第 2 学期）", size=14, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "题目：智能坐姿与学习环境监测系统设计", size=16, bold=True, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4):
        doc.add_paragraph()
    cover_items = [
        ("课程名称", "电子系统设计"),
        ("学    院", "________________"),
        ("专业班级", "________________"),
        ("学生姓名", "宋佳骏"),
        ("学    号", "________________"),
    ]
    for label, value in cover_items:
        add_para(doc, f"{label}：{value}", size=14, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    doc.add_page_break()

    add_para(doc, "智能坐姿与学习环境监测系统设计", size=18, bold=True, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_heading(doc, "摘要", 1)
    add_para(
        doc,
        "针对学生长时间伏案学习过程中容易出现前倾过度、环境光照不足等问题，本文设计了一种基于 STM32F103C8T6 的智能坐姿与学习环境监测系统。系统以 STM32F103C8T6 最小系统板为控制核心，采用 MPU6050 获取姿态信息，利用光敏电阻模块完成环境光照采集，并通过 OLED 显示当前坐姿角度、光照强度、报警角度阈值和报警延迟参数。系统设置旋转编码器作为人机交互输入，用户可在三状态菜单中调节坐姿报警角度和报警延迟时间。当检测到前倾角度超过设定阈值并持续超过设定时间时，低电平触发有源蜂鸣器报警；当环境光照 ADC 值低于设定阈值时，低电平点亮 LED 进行提醒。测试结果表明，该系统能够完成姿态检测、光照检测、参数设置和声光报警等预期功能，具有结构简单、成本低、交互直观和便于扩展等特点，符合电子系统设计课程对复杂电子系统综合设计的要求。",
    )
    add_para(doc, "关键词：STM32F103C8T6；MPU6050；光敏电阻；OLED；旋转编码器；坐姿监测", bold=True)

    add_heading(doc, "前言", 1)
    add_para(
        doc,
        "随着电子信息技术在学习、办公和健康监测场景中的应用不断增加，利用嵌入式系统完成局部环境感知和行为提醒已成为课程设计中较有代表性的综合应用方向。学生在长期学习时，坐姿不良和照明不足是比较常见的问题。前倾过度容易造成颈椎、腰椎疲劳，光照不足则可能加重视觉疲劳。传统提醒方式主要依靠人工经验判断，缺乏实时性和量化依据。因此，设计一种能够实时采集姿态和光照信息、并在本地完成显示与报警的电子系统，具有明确的实际意义。",
    )
    add_para(
        doc,
        "本文所设计系统属于嵌入式、传感器、模拟采集和数字控制混合应用系统。系统既包含 MPU6050 与 OLED 的 I2C 数字通信，也包含光敏电阻的 ADC 模拟量采集，还包含旋转编码器外部中断输入、蜂鸣器与 LED 输出控制。通过该课题可以综合训练 STM32 GPIO、ADC、外部中断、软件 I2C、状态机和模块化程序设计等内容。",
    )

    add_heading(doc, "① 系统的设计要求与技术指标的确定", 1)
    add_heading(doc, "1.1 设计任务与功能要求", 2)
    add_para(
        doc,
        "本系统面向学生学习场景，设计目标是对坐姿前倾状态和学习环境光照进行实时监测，并通过显示和声光方式向用户反馈。系统应在上电后完成 OLED、MPU6050、ADC、旋转编码器及报警输出端口初始化，随后进入主监控界面。主界面应实时显示坐姿角度和光照 ADC 数值，用户通过编码器按键进入设置界面，通过旋转编码器调节报警角度阈值和报警延迟时间。",
    )
    add_para(
        doc,
        "系统主要功能包括：姿态采集与角度估算、光照模拟量采集、OLED 实时显示、旋转编码器菜单设置、坐姿超限蜂鸣器报警、光照不足 LED 提醒以及三状态有限状态机控制。各功能之间既相互独立又由主控统一调度，能够体现一个完整电子系统从传感、处理、交互到执行输出的工作链路。",
    )
    add_heading(doc, "1.2 技术指标", 2)
    add_table(
        doc,
        "表 1 系统主要技术指标",
        ["项目", "设计指标"],
        [
            ["主控芯片", "STM32F103C8T6"],
            ["工作电压", "以 3.3 V 为主，部分低电平触发模块可按规格接 5 V"],
            ["姿态检测", "MPU6050 三轴加速度计和三轴陀螺仪，I2C 通信"],
            ["光照检测", "光敏电阻模块 AO 输出，PA0 / ADC1_IN0，12 位 ADC"],
            ["显示模块", "0.96 寸 I2C OLED，显示角度、光照、阈值和延迟"],
            ["交互输入", "360 度旋转编码器 A/B 相及 SW 按键"],
            ["角度显示范围", "0-90°"],
            ["报警角度范围", "0-90°，默认 25°"],
            ["报警延迟范围", "1-10 s，默认 3 s"],
            ["坐姿报警输出", "PB12，低电平触发有源蜂鸣器"],
            ["光照提醒输出", "PB13，低电平点亮警示 LED"],
            ["软件环境", "Keil MDK、STM32 标准库、C 语言"],
        ],
        widths_cm=[4.0, 11.0],
    )
    add_heading(doc, "1.3 系统约束与评价目标", 2)
    add_para(
        doc,
        "系统在设计上遵循低成本、易搭建、可调试和可演示原则。硬件模块均选用常见教学套件，便于学生在面包板上完成搭建；软件采用裸机循环加状态机结构，不引入操作系统，降低实现复杂度；人机交互采用 OLED 与旋转编码器，保证功能演示直观。评价重点包括系统功能完整性、参数设计合理性、程序结构清晰度、报警响应正确性以及硬件连接可靠性。",
    )

    add_heading(doc, "② 方案选择与可行性论证", 1)
    add_heading(doc, "2.1 总体方案", 2)
    add_para(
        doc,
        "系统总体结构如图 1 所示。左侧为信息采集模块，包括 MPU6050 姿态传感器、光敏电阻模块和旋转编码器；中间为 STM32F103C8T6 主控单元，负责数据采样、角度计算、阈值判断和状态机切换；右侧为显示与报警输出模块，包括 OLED、蜂鸣器和 LED。系统通过数字通信、模拟采集和 GPIO 控制共同完成学习环境监测任务。",
    )
    doc.add_picture(str(fig1), width=Cm(15.2))
    add_caption(doc, "图 1 系统总体架构与信号流程图")
    add_source(doc, "资料来源：作者根据系统功能划分绘制。")
    add_heading(doc, "2.2 主控芯片方案论证", 2)
    add_para(
        doc,
        "若采用 51 单片机，虽然成本较低、资料丰富，但 ADC、I2C 和多路中断扩展较繁琐，难以体现较完整的传感采集和菜单交互。若采用 Arduino UNO，开发门槛较低，但对寄存器和外设底层控制体现不足，与本课程 STM32 标准库学习内容匹配度不高。STM32F103C8T6 具备 ADC、GPIO、EXTI、定时器和丰富的接口资源，主频和运算能力能够满足角度计算、OLED 刷新和报警判断需求，且最小系统板成本低、资料多。因此本系统选择 STM32F103C8T6 作为主控芯片。",
    )
    add_heading(doc, "2.3 传感器与交互方案论证", 2)
    add_para(
        doc,
        "姿态检测方面，机械倾角开关只能输出开关量，无法显示连续角度，不适合可调阈值报警。MPU6050 可同时输出三轴加速度和角速度，使用 I2C 接口即可获得连续姿态数据，能够满足坐姿前倾角度估算需求。光照检测方面，数字光敏模块 DO 只能判断明暗，阈值依赖模块电位器；采用 AO 模拟输出并接入 STM32 ADC 后，可以显示 0-4095 范围内的连续光照值，更利于观察和软件阈值处理。交互方面，旋转编码器既能旋转调节数值，又带有按键，可用少量引脚完成菜单设置，比多个独立按键更节省硬件资源。",
    )
    add_heading(doc, "2.4 可行性分析", 2)
    add_para(
        doc,
        "硬件方面，OLED、MPU6050、光敏电阻、旋转编码器、蜂鸣器和 LED 均为常见模块，供电和接线方式清晰，适合课程实验环境搭建。软件方面，底层驱动基于 STM32 标准库和课程例程，应用层主要完成状态机、阈值判断和数据显示，难度适中。成本方面，系统由低成本模块组成，不需要复杂专用仪器。调试方面，OLED 可显示关键实时数据，便于分模块验证。综上，系统方案在功能、成本、实现难度和课程契合度上均具有可行性。",
    )

    add_heading(doc, "③ 单元电路的设计、参数计算和元器件选择", 1)
    add_heading(doc, "3.1 主控最小系统与引脚分配", 2)
    add_para(
        doc,
        "主控采用 STM32F103C8T6 最小系统板。该芯片基于 ARM Cortex-M3 内核，具有较丰富的 GPIO、ADC 和中断资源。系统通过 ST-Link 的 SWDIO 与 SWCLK 接口进行程序下载，应用模块按功能分配到不同端口，避免 OLED 与 MPU6050 共用同一组软件 I2C 引脚造成驱动冲突。",
    )
    add_table(
        doc,
        "表 2 系统引脚分配表",
        ["功能模块", "模块引脚", "STM32 引脚", "说明"],
        [
            ["OLED", "SCL/SDA", "PB8/PB9", "OLED 软件 I2C"],
            ["MPU6050", "SCL/SDA", "PB10/PB11", "姿态传感器软件 I2C"],
            ["光敏电阻", "AO", "PA0", "ADC1_IN0 模拟输入"],
            ["旋转编码器", "A/B", "PA1/PA2", "外部中断输入"],
            ["旋转编码器", "SW", "PA3", "按键输入，软件消抖"],
            ["蜂鸣器", "IN", "PB12", "低电平触发报警"],
            ["警示 LED", "IN", "PB13", "低电平点亮"],
            ["ST-Link", "SWDIO/SWCLK", "SWDIO/SWCLK", "程序下载与调试"],
        ],
        widths_cm=[3.0, 3.0, 3.0, 6.0],
    )
    doc.add_picture(str(fig3), width=Cm(15.2))
    add_caption(doc, "图 2 硬件模块连接与引脚分配示意图")
    add_source(doc, "资料来源：作者根据实物接线关系绘制。")

    add_heading(doc, "3.2 OLED 显示单元", 2)
    add_para(
        doc,
        "OLED 采用 0.96 寸 I2C 显示屏，常见引脚为 GND、VCC、SCL、SDA。本系统将 OLED 的 SCL 接 PB8、SDA 接 PB9，使用软件 I2C 驱动。主监控界面显示系统名称、当前前倾角度、光照 ADC 值、设定角度阈值和报警延迟时间；设置界面分别显示角度阈值和延迟时间。OLED 模块显示清晰、功耗较低、占用引脚少，适合小型课程设计系统。",
    )
    add_heading(doc, "3.3 MPU6050 姿态检测单元", 2)
    add_para(
        doc,
        "MPU6050 模块集成三轴加速度计和三轴陀螺仪，可通过 I2C 总线读取姿态数据。系统将 MPU6050 的 SCL 接 PB10、SDA 接 PB11，与 OLED 分离成两组软件 I2C，降低调试复杂度。MPU6050 常用地址与 AD0 引脚电平有关，AD0 为低时写地址为 0xD0，AD0 为高时写地址为 0xD2。程序中自动尝试两种地址，提高了不同模块批次下的兼容性。",
    )
    add_para(
        doc,
        "坐姿前倾角度主要由三轴加速度数据估算。设 AccX 表示安装方向下对前倾变化敏感的轴，AccY 和 AccZ 表示另外两个正交方向，则前倾角可按式（1）估算：",
    )
    add_para(doc, "Angle = atan(|AccX| / sqrt(AccY^2 + AccZ^2)) × 180 / π    （1）", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para(
        doc,
        "该算法不依赖陀螺仪积分，计算量小，适合裸机循环实时运行。若实际安装方向不同，只需将敏感轴由 AccX 调整为 AccY 或 AccZ。由于本系统主要用于课程演示和趋势提醒，采用加速度估算角度能够满足设计需求。",
    )
    add_heading(doc, "3.4 光照检测与 ADC 参数计算", 2)
    add_para(
        doc,
        "光敏检测单元采用带 AO 输出的光敏电阻模块。模块输出模拟电压接入 PA0，即 ADC1_IN0。STM32F103C8T6 的 ADC 为 12 位，数字量范围为 0-4095。当参考电压为 3.3 V 时，输入电压与 ADC 数值关系如式（2）所示：",
    )
    add_para(doc, "Vin = ADC_Value / 4095 × 3.3 V    （2）", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para(
        doc,
        "例如 ADC_Value 为 2048 时，输入电压约为 1.65 V。程序中将光照报警阈值设置为 1200，当采样值低于该阈值时认为环境光照不足，点亮警示 LED。实际使用中，阈值可根据光敏模块输出方向和环境亮度通过实验进行微调。",
    )
    add_heading(doc, "3.5 旋转编码器输入单元", 2)
    add_para(
        doc,
        "旋转编码器用于参数设置。A 相接 PA1，B 相接 PA2，配置为上拉输入并使用外部中断检测下降沿，通过另一相电平判断旋转方向。SW 按键接 PA3，程序采用软件消抖方式检测按下事件。在主监控界面按下 SW 进入角度设置状态，在角度设置状态按下进入延迟设置状态，在延迟设置状态按下返回主监控界面。该方式只需三个输入引脚即可完成菜单切换和数值调节。",
    )
    add_heading(doc, "3.6 声光报警输出单元", 2)
    add_para(
        doc,
        "蜂鸣器采用低电平触发有源蜂鸣器模块，IN 接 PB12。当坐姿角度连续超过设定阈值，并且持续时间超过 Alarm_Delay 后，程序将 PB12 拉低，蜂鸣器发声；否则 PB12 输出高电平关闭蜂鸣器。LED 采用低电平点亮模块，IN 接 PB13。当光照 ADC 值低于阈值时，PB13 拉低点亮 LED，提示当前环境光照不足。",
    )
    add_para(
        doc,
        "若使用普通 LED 而非模块，需要串联限流电阻。设供电电压为 3.3 V，LED 正向压降约为 2.0 V，期望电流取 5 mA，则限流电阻可按式（3）计算：",
    )
    add_para(doc, "R = (3.3 - 2.0) / 0.005 = 260 Ω    （3）", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para(doc, "实际选型可取常见阻值 330 Ω 或 470 Ω，以保证 LED 有足够亮度并限制单片机引脚电流。")

    add_heading(doc, "3.7 软件状态机设计", 2)
    add_para(
        doc,
        "系统软件采用三状态有限状态机。STATE_MONITOR 为主监控界面，周期性采集 MPU6050 与 ADC 数据，计算角度并执行报警判断；STATE_SET_ANGLE 为角度阈值设置界面，编码器旋转改变 Alarm_Angle，范围为 0-90°；STATE_SET_DELAY 为报警延迟设置界面，编码器旋转改变 Alarm_Delay，范围为 1-10 s。状态转换由 PA3 按键触发，流程如图 3 所示。",
    )
    doc.add_picture(str(fig2), width=Cm(15.2))
    add_caption(doc, "图 3 软件三状态菜单流程图")
    add_source(doc, "资料来源：作者根据主循环状态机程序绘制。")

    add_heading(doc, "3.8 程序关键参数与报警逻辑", 2)
    add_table(
        doc,
        "表 3 程序关键参数",
        ["参数名", "默认值", "含义"],
        [
            ["LIGHT_ALARM_THRESHOLD", "1200", "光照 ADC 报警阈值，低于该值点亮 LED"],
            ["SAMPLE_PERIOD_MS", "100 ms", "主循环采样周期"],
            ["Alarm_Angle", "25°", "默认坐姿报警角度阈值"],
            ["Alarm_Delay", "3 s", "默认坐姿持续超限报警延迟"],
            ["ALARM_ANGLE_MAX", "90°", "角度阈值上限"],
            ["ALARM_DELAY_MAX", "10 s", "报警延迟上限"],
        ],
        widths_cm=[5.0, 3.0, 7.0],
    )
    add_para(
        doc,
        "报警延迟通过计数方式实现。由于采样周期为 100 ms，若 Alarm_Delay 为 3 s，则需要连续超限 30 次采样后才触发蜂鸣器。该设计能够过滤短时间姿态晃动，避免用户轻微动作导致频繁误报警。若角度恢复到阈值以下，超限计数清零，蜂鸣器关闭。",
    )
    add_heading(doc, "3.9 系统测试与误差分析", 2)
    add_para(
        doc,
        "系统调试按从单元到整体的顺序进行。首先验证 OLED 是否能够正常显示启动信息；随后连接 MPU6050，观察角度随模块倾斜而变化；再连接光敏模块，遮挡或照射光敏电阻，确认 Light 数值变化；之后验证旋转编码器能否进入设置界面并调节参数；最后验证蜂鸣器和 LED 的报警逻辑。实际测试表明，系统能够完成预期功能，OLED 显示稳定，参数调节有效，蜂鸣器和 LED 输出符合低电平有效逻辑。",
    )
    add_para(
        doc,
        "系统误差主要来自三个方面。第一，MPU6050 安装方向和固定方式会影响前倾角度计算，模块若松动会造成角度波动。第二，光敏电阻模块输出与实际照度并非严格线性，ADC 数值只能作为相对光照强度使用。第三，机械编码器存在触点抖动，快速旋转时可能出现计数跳变。针对这些问题，可通过固定传感器、增加数字滤波、平均多次 ADC 采样、优化编码器消抖或引入定时器编码器模式等方式改进。",
    )

    add_heading(doc, "结论", 1)
    add_para(
        doc,
        "本文完成了一种基于 STM32F103C8T6 的智能坐姿与学习环境监测系统设计。系统综合使用 MPU6050、光敏电阻、OLED、旋转编码器、蜂鸣器和 LED 等模块，实现了姿态角度计算、光照强度采集、OLED 实时显示、参数菜单设置以及声光报警功能。系统硬件连接清晰，软件采用三状态有限状态机组织，逻辑结构明确，便于调试和功能扩展。",
    )
    add_para(
        doc,
        "从课程设计角度看，该系统覆盖了嵌入式主控、数字通信、模拟采集、外部中断、GPIO 输出和状态机程序设计等多个知识点，满足《电子系统设计》课程对复杂电子系统综合设计的要求。后续可进一步加入数据存储、蓝牙或 Wi-Fi 远程传输、蜂鸣器分级提示、光照标定以及更复杂的姿态融合算法，以提高系统实用性和测量稳定性。",
    )

    add_heading(doc, "④ 参考资料目录", 1)
    refs = [
        "STMicroelectronics. STM32F103x8/B Datasheet[M]. Geneva: STMicroelectronics, 2015.",
        "STMicroelectronics. RM0008 Reference Manual: STM32F10xxx Advanced ARM-based 32-bit MCUs[M]. Geneva: STMicroelectronics, 2021.",
        "STMicroelectronics. STM32F10x Standard Peripheral Library User Manual[M]. Geneva: STMicroelectronics, 2011.",
        "InvenSense. MPU-6000 and MPU-6050 Product Specification[M]. San Jose: InvenSense, 2013.",
        "InvenSense. MPU-6000 and MPU-6050 Register Map and Descriptions[M]. San Jose: InvenSense, 2013.",
        "Solomon Systech. SSD1306 OLED Driver Controller Datasheet[M]. Hong Kong: Solomon Systech, 2008.",
        "江科大. STM32 入门教程及配套标准库例程[Z].",
        "谭浩强. C 程序设计[M]. 北京: 清华大学出版社, 2017.",
        "童诗白, 华成英. 模拟电子技术基础[M]. 北京: 高等教育出版社, 2015.",
        "王田苗. 嵌入式系统设计与实践[M]. 北京: 清华大学出版社, 2019.",
    ]
    for i, ref in enumerate(refs, 1):
        add_para(doc, f"[{i}] {ref}", first_line=False, after=3)

    add_heading(doc, "附录：核心程序与工程说明", 1)
    add_para(
        doc,
        "本设计的 Keil 工程位于 Project 文件夹，主程序文件为 Project\\User\\main.c。底层驱动包括 OLED.c、MPU6050.c、MyI2C.c、AD.c 和 Encoder.c。工程已使用 ARM Compiler 6 完成编译，生成 Project.axf 和 Project.hex，可通过 ST-Link 的 SWD 接口下载到 STM32F103C8T6 最小系统板运行。",
    )

    # Footer page number
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        r._r.append(fldChar1)
        r._r.append(instrText)
        r._r.append(fldChar2)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    output = create_document()
    print(output)
